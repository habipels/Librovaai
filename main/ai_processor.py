"""
AI Entegrasyonu - Dosya İşleme ve İçerik Analizi
PDF ve Word dosyalarını işleyip içindekiler çıkarır, özet üretir
"""

import os
import re
from typing import Dict, List, Tuple, Optional
from django.core.files.uploadedfile import UploadedFile
import logging

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Doküman işleme ana sınıfı"""
    
    def __init__(self, file_path: str = None, uploaded_file: UploadedFile = None):
        self.file_path = file_path
        self.uploaded_file = uploaded_file
        self.content = ""
        self.file_type = None
        
        if file_path:
            self.file_type = self._detect_file_type(file_path)
        elif uploaded_file:
            self.file_type = self._detect_file_type(uploaded_file.name)
    
    def _detect_file_type(self, filename: str) -> str:
        """Dosya tipini belirle"""
        ext = filename.lower().split('.')[-1]
        if ext == 'pdf':
            return 'pdf'
        elif ext in ['docx', 'doc']:
            return 'docx'
        else:
            raise ValueError(f"Desteklenmeyen dosya tipi: {ext}")
    
    def extract_text(self) -> str:
        """Dosyadan metin çıkar"""
        if self.file_type == 'pdf':
            return self._extract_from_pdf()
        elif self.file_type == 'docx':
            return self._extract_from_docx()
        else:
            raise ValueError("Desteklenmeyen dosya tipi")
    
    def _extract_from_pdf(self) -> str:
        """PDF'den metin çıkar"""
        try:
            import PyPDF2
            
            text = ""
            
            if self.file_path:
                with open(self.file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
            
            elif self.uploaded_file:
                pdf_reader = PyPDF2.PdfReader(self.uploaded_file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            
            self.content = text
            return text
            
        except ImportError:
            logger.error("PyPDF2 kütüphanesi yüklü değil. pip install PyPDF2")
            raise ImportError("PyPDF2 kütüphanesi gerekli: pip install PyPDF2")
        except Exception as e:
            logger.error(f"PDF okuma hatası: {str(e)}")
            raise Exception(f"PDF okuma hatası: {str(e)}")
    
    def _extract_from_docx(self) -> str:
        """Word dosyasından metin çıkar"""
        try:
            import docx
            
            text = ""
            
            if self.file_path:
                doc = docx.Document(self.file_path)
            elif self.uploaded_file:
                doc = docx.Document(self.uploaded_file)
            else:
                raise ValueError("Dosya bulunamadı")
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Tabloları da işle
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            self.content = text
            return text
            
        except ImportError:
            logger.error("python-docx kütüphanesi yüklü değil. pip install python-docx")
            raise ImportError("python-docx kütüphanesi gerekli: pip install python-docx")
        except Exception as e:
            logger.error(f"Word okuma hatası: {str(e)}")
            raise Exception(f"Word okuma hatası: {str(e)}")
    
    def extract_table_of_contents(self) -> List[Dict]:
        """İçindekiler çıkar"""
        if not self.content:
            self.extract_text()
        
        toc = []
        
        # Yaygın başlık kalıpları
        patterns = [
            # "1. Başlık", "1.1 Alt Başlık"
            r'^(\d+(?:\.\d+)*)\s+(.+)$',
            # "Bölüm 1", "Bölüm 1.1"
            r'^(?:Bölüm|BÖLÜM|Chapter|CHAPTER)\s+(\d+(?:\.\d+)*)[:\.]?\s*(.+)$',
            # "I. Başlık", "II. Başlık" (Roma rakamları)
            r'^([IVXLCDM]+)\.\s+(.+)$',
            # Büyük harfle yazılmış başlıklar (en az 3 kelime)
            r'^([A-ZÇĞİÖŞÜ\s]{10,})$',
        ]
        
        lines = self.content.split('\n')
        
        for i, line in enumerate(lines):
            line = line.strip()
            
            if not line or len(line) < 3:
                continue
            
            for pattern in patterns:
                match = re.match(pattern, line)
                if match:
                    if len(match.groups()) == 2:
                        number, title = match.groups()
                        level = self._calculate_level(number)
                    else:
                        # Büyük harfle yazılmış başlıklar için
                        title = match.group(1) if len(match.groups()) == 1 else match.group(2)
                        number = str(len(toc) + 1)
                        level = 1
                    
                    toc.append({
                        'number': number.strip(),
                        'title': title.strip(),
                        'level': level,
                        'line_number': i,
                    })
                    break
        
        return toc
    
    def _calculate_level(self, number: str) -> int:
        """Başlık seviyesini hesapla (1.1.1 -> level 3)"""
        if '.' in number:
            return len(number.split('.'))
        return 1
    
    def extract_chapters_with_content(self, max_chapters: int = 50) -> List[Dict]:
        """Bölümleri içerikleriyle birlikte çıkar"""
        if not self.content:
            self.extract_text()
        
        toc = self.extract_table_of_contents()
        
        if not toc:
            # İçindekiler bulunamadıysa, metni bölümlere ayır
            return self._split_content_into_chapters()
        
        lines = self.content.split('\n')
        chapters = []
        
        for i, chapter in enumerate(toc[:max_chapters]):
            start_line = chapter['line_number']
            
            # Sonraki bölümün başlangıcını bul
            if i + 1 < len(toc):
                end_line = toc[i + 1]['line_number']
            else:
                end_line = len(lines)
            
            # Bölüm içeriğini al
            chapter_content = '\n'.join(lines[start_line:end_line]).strip()
            
            chapters.append({
                'number': chapter['number'],
                'title': chapter['title'],
                'level': chapter['level'],
                'content': chapter_content[:5000],  # İlk 5000 karakter
                'full_length': len(chapter_content)
            })
        
        return chapters
    
    def _split_content_into_chapters(self) -> List[Dict]:
        """İçindekiler bulunamazsa metni bölümlere ayır"""
        # Metni paragraf paragraf ayır
        paragraphs = [p.strip() for p in self.content.split('\n\n') if p.strip()]
        
        # Her 500 kelimelik parçaları bölüm olarak kabul et
        chapters = []
        current_chapter = []
        word_count = 0
        chapter_num = 1
        
        for para in paragraphs:
            words = len(para.split())
            
            if word_count + words > 500 and current_chapter:
                # Yeni bölüm oluştur
                chapters.append({
                    'number': str(chapter_num),
                    'title': f'Bölüm {chapter_num}',
                    'level': 1,
                    'content': '\n\n'.join(current_chapter),
                    'full_length': len('\n\n'.join(current_chapter))
                })
                
                current_chapter = [para]
                word_count = words
                chapter_num += 1
            else:
                current_chapter.append(para)
                word_count += words
        
        # Son bölümü ekle
        if current_chapter:
            chapters.append({
                'number': str(chapter_num),
                'title': f'Bölüm {chapter_num}',
                'level': 1,
                'content': '\n\n'.join(current_chapter),
                'full_length': len('\n\n'.join(current_chapter))
            })
        
        return chapters


class AIContentGenerator:
    """AI ile içerik üretimi (OpenAI, Gemini veya lokal model)"""
    
    def __init__(self, api_key: str = None, model: str = "gpt-3.5-turbo"):
        self.api_key = api_key
        self.model = model
        self.use_ai = bool(api_key)
    
    def generate_book_summary(self, content: str, max_length: int = 500) -> str:
        """Kitap özeti üret"""
        
        if self.use_ai:
            return self._generate_ai_summary(content, max_length)
        else:
            return self._generate_simple_summary(content, max_length)
    
    def generate_chapter_summary(self, chapter_content: str, max_length: int = 200) -> str:
        """Bölüm özeti üret"""
        
        if self.use_ai:
            return self._generate_ai_summary(chapter_content, max_length, is_chapter=True)
        else:
            return self._generate_simple_summary(chapter_content, max_length)
    
    def _generate_ai_summary(self, content: str, max_length: int, is_chapter: bool = False) -> str:
        """AI ile özet üret (OpenAI API)"""
        try:
            import openai
            
            openai.api_key = self.api_key
            
            content_type = "bölüm" if is_chapter else "kitap"
            
            prompt = f"""
            Aşağıdaki {content_type} içeriğini Türkçe olarak özetle. 
            Özet en fazla {max_length} kelime olsun ve ana fikirleri kapsasın.
            
            İçerik:
            {content[:3000]}
            
            Özet:
            """
            
            response = openai.ChatCompletion.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "Sen bir kitap özeti uzmanısın."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=max_length * 2,
                temperature=0.7
            )
            
            summary = response.choices[0].message.content.strip()
            return summary
            
        except ImportError:
            logger.warning("OpenAI kütüphanesi yüklü değil, basit özet kullanılıyor")
            return self._generate_simple_summary(content, max_length)
        except Exception as e:
            logger.error(f"AI özet üretme hatası: {str(e)}")
            return self._generate_simple_summary(content, max_length)
    
    def _generate_simple_summary(self, content: str, max_length: int = 500) -> str:
        """Basit özet üret (AI olmadan)"""
        # İlk paragrafları al
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
        
        summary = ""
        word_count = 0
        
        for para in paragraphs[:5]:  # İlk 5 paragraf
            words = para.split()
            if word_count + len(words) <= max_length:
                summary += para + "\n\n"
                word_count += len(words)
            else:
                # Kalan kelime limitini kullan
                remaining = max_length - word_count
                summary += ' '.join(words[:remaining]) + "..."
                break
        
        return summary.strip() or "İçerik özeti oluşturulamadı."


def process_book_file(book_instance, use_ai: bool = False, api_key: str = None):
    """
    Kitap dosyasını işle, bölümlere ayır, özetleri oluştur
    
    Args:
        book_instance: Book model instance
        use_ai: AI kullanılsın mı?
        api_key: OpenAI API key (opsiyonel)
    
    Returns:
        dict: İşlem sonucu
    """
    from main.models import Chapter
    
    try:
        # Dosya işleyiciyi oluştur
        processor = DocumentProcessor(file_path=book_instance.file.path)
        
        # Metni çıkar
        content = processor.extract_text()
        
        # Bölümleri çıkar
        chapters_data = processor.extract_chapters_with_content()
        
        # AI oluşturucu
        ai_generator = AIContentGenerator(api_key=api_key)
        
        # Kitap özeti oluştur
        if use_ai or not book_instance.ai_summary:
            book_summary = ai_generator.generate_book_summary(content[:5000])
            book_instance.ai_summary = book_summary
        
        # Bölümleri veritabanına kaydet
        created_chapters = []
        for i, chapter_data in enumerate(chapters_data):
            
            # Bölüm özeti oluştur
            chapter_summary = ai_generator.generate_chapter_summary(
                chapter_data['content']
            )
            
            chapter, created = Chapter.objects.update_or_create(
                book=book_instance,
                chapter_number=i + 1,
                defaults={
                    'title': chapter_data['title'],
                    'content': chapter_data['content'],
                    'level': chapter_data['level'],
                    'order': i,
                    'ai_summary': chapter_summary
                }
            )
            
            created_chapters.append(chapter)
        
        # Kitap bilgilerini güncelle
        book_instance.is_processed = True
        book_instance.has_toc = len(chapters_data) > 0
        book_instance.save()
        
        return {
            'success': True,
            'chapters_count': len(created_chapters),
            'has_summary': bool(book_instance.ai_summary),
            'message': f'{len(created_chapters)} bölüm başarıyla işlendi.'
        }
        
    except Exception as e:
        logger.error(f"Kitap işleme hatası: {str(e)}")
        return {
            'success': False,
            'error': str(e),
            'message': f'Dosya işlenirken hata oluştu: {str(e)}'
        }
