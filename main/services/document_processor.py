"""
Dosya İşleme Servisleri
PDF ve Word dosyalarını okuma, içindekiler çıkarma, bölümlere ayırma
"""
import os
import re
from typing import List, Dict, Tuple
from django.core.files.uploadedfile import UploadedFile


class DocumentProcessor:
    """
    PDF ve Word belgelerini işleyen ana sınıf
    """
    
    @staticmethod
    def get_file_type(file: UploadedFile) -> str:
        """Dosya tipini belirler"""
        ext = file.name.split('.')[-1].lower()
        return ext
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """
        PDF dosyasından metin çıkarır
        Gerekli: pip install PyPDF2
        """
        try:
            import PyPDF2
            
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text
        except ImportError:
            return "PyPDF2 yüklü değil. Lütfen: pip install PyPDF2"
        except Exception as e:
            return f"PDF okuma hatası: {str(e)}"
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """
        Word (docx) dosyasından metin çıkarır
        Gerekli: pip install python-docx
        """
        try:
            from docx import Document
            
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except ImportError:
            return "python-docx yüklü değil. Lütfen: pip install python-docx"
        except Exception as e:
            return f"DOCX okuma hatası: {str(e)}"
    
    @staticmethod
    def extract_text_from_doc(file_path: str) -> str:
        """
        Eski Word (doc) dosyasından metin çıkarır
        Gerekli: pip install textract veya antiword
        """
        try:
            import textract
            text = textract.process(file_path).decode('utf-8')
            return text
        except ImportError:
            return "textract yüklü değil. Lütfen: pip install textract"
        except Exception as e:
            return f"DOC okuma hatası: {str(e)}"
    
    def extract_text(self, file_path: str, file_type: str) -> str:
        """Dosya tipine göre metin çıkarır"""
        if file_type == 'pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_type == 'docx':
            return self.extract_text_from_docx(file_path)
        elif file_type == 'doc':
            return self.extract_text_from_doc(file_path)
        else:
            return f"Desteklenmeyen dosya tipi: {file_type}"


class TableOfContentsExtractor:
    """
    Belgeden içindekiler (Table of Contents) çıkarır
    """
    
    @staticmethod
    def extract_toc_patterns(text: str) -> List[Dict[str, any]]:
        """
        Metin içinden başlıkları tespit eder
        Şu kalıpları arar:
        - Bölüm 1, Bölüm 2, ...
        - BÖLÜM I, BÖLÜM II, ...
        - 1. Başlık, 2. Başlık, ...
        - Büyük harfle yazılmış başlıklar
        """
        chapters = []
        lines = text.split('\n')
        
        # Kalıplar
        patterns = [
            r'^(BÖLÜM|Bölüm|CHAPTER|Chapter)\s+(\d+|[IVXLCDM]+)[:\.\s]+(.+)$',
            r'^(\d+)\.\s+(.+)$',
            r'^([A-Z][A-ZÜĞŞIÖÇ\s]{10,})$',  # Büyük harfle yazılmış uzun satırlar
        ]
        
        order = 0
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
            
            for pattern in patterns:
                match = re.match(pattern, line)
                if match:
                    order += 1
                    chapters.append({
                        'order': order,
                        'title': line,
                        'line_number': i,
                        'level': 1,
                    })
                    break
        
        return chapters
    
    @staticmethod
    def extract_toc_from_docx(file_path: str) -> List[Dict[str, any]]:
        """
        Word belgesinden heading stilleri kullanarak içindekiler çıkarır
        """
        try:
            from docx import Document
            
            doc = Document(file_path)
            chapters = []
            order = 0
            
            for i, paragraph in enumerate(doc.paragraphs):
                # Heading stillerini kontrol et
                if paragraph.style.name.startswith('Heading'):
                    level = int(paragraph.style.name[-1]) if paragraph.style.name[-1].isdigit() else 1
                    order += 1
                    chapters.append({
                        'order': order,
                        'title': paragraph.text,
                        'line_number': i,
                        'level': level,
                    })
            
            return chapters
        except Exception as e:
            print(f"DOCX TOC çıkarma hatası: {e}")
            return []


class ChapterExtractor:
    """
    Belgeyi bölümlere ayırır
    """
    
    @staticmethod
    def split_into_chapters(text: str, toc: List[Dict[str, any]]) -> List[Dict[str, any]]:
        """
        İçindekiler bilgisini kullanarak metni bölümlere ayırır
        """
        if not toc:
            # Eğer içindekiler yoksa, tüm metni tek bölüm olarak döndür
            return [{
                'order': 1,
                'title': 'Tüm İçerik',
                'content': text,
                'level': 1,
            }]
        
        lines = text.split('\n')
        chapters = []
        
        for i, chapter_info in enumerate(toc):
            start_line = chapter_info['line_number']
            end_line = toc[i + 1]['line_number'] if i + 1 < len(toc) else len(lines)
            
            chapter_content = '\n'.join(lines[start_line:end_line])
            
            chapters.append({
                'order': chapter_info['order'],
                'title': chapter_info['title'],
                'content': chapter_content,
                'level': chapter_info.get('level', 1),
            })
        
        return chapters


# Kullanım örneği
def process_book_file(file_path: str, file_type: str) -> Tuple[str, List[Dict], List[Dict]]:
    """
    Kitap dosyasını işler ve metin, içindekiler ve bölümleri döndürür
    
    Returns:
        Tuple[str, List[Dict], List[Dict]]: (metin, içindekiler, bölümler)
    """
    processor = DocumentProcessor()
    toc_extractor = TableOfContentsExtractor()
    chapter_extractor = ChapterExtractor()
    
    # 1. Metni çıkar
    text = processor.extract_text(file_path, file_type)
    
    # 2. İçindekiler çıkar
    if file_type == 'docx':
        toc = toc_extractor.extract_toc_from_docx(file_path)
    else:
        toc = toc_extractor.extract_toc_patterns(text)
    
    # 3. Bölümlere ayır
    chapters = chapter_extractor.split_into_chapters(text, toc)
    
    return text, toc, chapters
